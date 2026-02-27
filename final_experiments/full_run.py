#!/usr/bin/env python3
import os
import re
import random
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
from tqdm import tqdm
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
from dotenv import load_dotenv
from docx import Document

# --- Gemini SDK (new) ---
from google import genai


# =========================
# Config
# =========================

RUBRIC = """You are an instructor grading different essays.

Grade holistically using these dimensions:
1) Flow (clarity/readability across sentences)
2) Transitions (connections between ideas/paragraphs)
3) Content Quality & Focus (relevance, specificity, alignment to the week’s concepts)
4) Spelling & Grammar (correctness, professionalism)
5) Knowledge & Depth (understanding + thoughtful engagement)
6) Structure (organization, paragraphing, logical progression)

Score rules:
- Provide ONE overall score from 1 to 10 (integers only).
- 1 = fundamentally flawed; minimal understanding; very poor writing.
- 5 = mixed/adequate; some understanding; clear weaknesses.
- 10 = exemplary; polished, insightful, well-structured; no meaningful weaknesses.
- Avoid score inflation: 9–10 only if nearly flawless.
"""

JSON_CONTRACT = """Output STRICT JSON ONLY (no markdown, no extra text).

Return a JSON array with EXACTLY one object per essay, in the SAME ORDER as the essays are presented.

Each object must match:
{
  "id": "<string>",
  "pred_score": <integer 1-10>,
  "justification": "<1 sentence>"
}
"""


@dataclass
class PathsConfig:
    base_dir: Path
    average_dir: Path
    exceptional_dir: Path
    poor_dir: Path
    outputs_dir: Path


class GeminiJSONError(Exception):
    pass


# =========================
# IO: DOCX → text
# =========================

def docx_to_text(docx_path: Path) -> str:
    """
    Extract text from .docx (paragraphs + tables).
    """
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX not found: {docx_path}")

    doc = Document(str(docx_path))
    parts: List[str] = []

    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            row_text: List[str] = []
            for cell in row.cells:
                ct = (cell.text or "").strip()
                if ct:
                    row_text.append(ct)
            if row_text:
                parts.append(" | ".join(row_text))

    return "\n\n".join(parts).strip()


def find_docx_by_stem(folder: Path, stem: str) -> Path:
    stem_lower = stem.lower()
    matches = [p for p in folder.glob("*.docx") if p.stem.lower() == stem_lower]
    if not matches:
        raise FileNotFoundError(f"Could not find {stem}.docx in {folder}")
    if len(matches) > 1:
        raise RuntimeError(f"Multiple matches for {stem}.docx in {folder}: {matches}")
    return matches[0]


def load_paths(base_dir_str: str) -> PathsConfig:
    base_dir = Path(base_dir_str).expanduser().resolve()
    average_dir = base_dir / "Average Essays"
    exceptional_dir = base_dir / "Exceptional Essays"
    poor_dir = base_dir / "Poor Essays"
    outputs_dir = Path(__file__).parent / "results"
    outputs_dir.mkdir(parents=True, exist_ok=True)

    for d in [average_dir, exceptional_dir, poor_dir]:
        if not d.exists():
            raise FileNotFoundError(f"Folder not found: {d}")

    return PathsConfig(
        base_dir=base_dir,
        average_dir=average_dir,
        exceptional_dir=exceptional_dir,
        poor_dir=poor_dir,
        outputs_dir=outputs_dir,
    )


# =========================
# Prompt builders
# =========================

def build_fewshot_prompt_multi(anchor_exceptional: str, anchor_average: str, anchor_poor: str,
                               essays: List[Tuple[str, str]]) -> str:
    essays_block = []
    for essay_id, essay_text in essays:
        essays_block.append(f"[ESSAY]\nID: {essay_id}\nTEXT:\n{essay_text}\n")
    essays_block = "\n".join(essays_block)

    return f"""{RUBRIC}

Calibration anchors (use these to calibrate your scoring scale):

[ANCHOR A — 10/10 EXAMPLE]
{anchor_exceptional}

[ANCHOR B — 5/10 EXAMPLE]
{anchor_average}

[ANCHOR C — 1/10 EXAMPLE]
{anchor_poor}

Now grade EACH essay below independently.

IMPORTANT RULES:
- Return a JSON ARRAY only.
- The array must have EXACTLY {len(essays)} items.
- Items must be in the SAME ORDER as the essays appear below.
- Use the essay ID exactly as provided.
- pred_score must be an integer 1–10.

Essays (in order):
{essays_block}

{JSON_CONTRACT}
"""


def build_zeroshot_prompt_multi(essays: List[Tuple[str, str]]) -> str:
    essays_block = []
    for essay_id, essay_text in essays:
        essays_block.append(f"[ESSAY]\nID: {essay_id}\nTEXT:\n{essay_text}\n")
    essays_block = "\n".join(essays_block)

    return f"""{RUBRIC}

Now grade EACH essay below independently.

IMPORTANT RULES:
- Return a JSON ARRAY only.
- The array must have EXACTLY {len(essays)} items.
- Items must be in the SAME ORDER as the essays appear below.
- Use the essay ID exactly as provided.
- pred_score must be an integer 1–10.

Essays (in order):
{essays_block}

{JSON_CONTRACT}
"""


# =========================
# Gemini call + JSON parsing
# =========================

def extract_json_array(text: str) -> str:
    text = (text or "").strip()
    # Strip markdown: ```json ... ``` or truncated ```json\n... (no closing ```)
    if "```" in text:
        m = re.search(r"```(?:json)?\s*([\s\S]*?)```", text)
        if m:
            text = m.group(1).strip()
        elif text.startswith("```"):
            text = re.sub(r"^```(?:json)?\s*\n?", "", text).strip()
    # Drop any leading/trailing prose: take from first '[' to last ']'
    start = text.find("[")
    if start != -1:
        end = text.rfind("]")
        if end != -1 and end >= start:
            text = text[start : end + 1]
    try:
        parsed = json.loads(text)
        if isinstance(parsed, list):
            return text
        if isinstance(parsed, dict):
            for key in ("results", "data", "items", "grades", "scores", "essays"):
                val = parsed.get(key)
                if isinstance(val, list):
                    return json.dumps(val, ensure_ascii=False)
    except Exception:
        pass
    if text.startswith("[") and text.endswith("]"):
        return text
    m = re.search(r"\[[\s\S]*\]", text)
    if m:
        return m.group(0).strip()
    raise GeminiJSONError("Could not find JSON array in model output.")


def validate_item(obj: Any, expected_id: str) -> Dict[str, Any]:
    if not isinstance(obj, dict):
        raise GeminiJSONError(f"Item is not an object: {obj}")

    got_id = str(obj.get("id", "")).strip()
    if got_id != expected_id:
        # Strict: enforce same-order / same-id
        raise GeminiJSONError(f"Expected id={expected_id}, got id={got_id}")

    score = obj.get("pred_score", None)
    if not isinstance(score, int):
        # try convert from string digits
        try:
            score = int(str(score).strip())
        except Exception:
            raise GeminiJSONError(f"pred_score not int-like for id={expected_id}: {obj.get('pred_score')}")

    clamped = False
    if score < 1:
        score = 1
        clamped = True
    if score > 10:
        score = 10
        clamped = True

    justification = str(obj.get("justification", "")).strip()
    if not justification:
        justification = ""

    return {
        "id": expected_id,
        "pred_score": score,
        "justification": justification,
        "clamped": clamped,
    }


def validate_list(arr: Any, expected_ids: List[str]) -> List[Dict[str, Any]]:
    if not isinstance(arr, list):
        raise GeminiJSONError(f"Expected list, got {type(arr)}")
    if len(arr) != len(expected_ids):
        raise GeminiJSONError(f"Expected {len(expected_ids)} items, got {len(arr)}")

    out = []
    for obj, exp_id in zip(arr, expected_ids):
        out.append(validate_item(obj, exp_id))
    return out


@retry(
    reraise=True,
    stop=stop_after_attempt(5),
    wait=wait_exponential(multiplier=1, min=2, max=20),
    retry=retry_if_exception_type((GeminiJSONError, RuntimeError, TimeoutError, ConnectionError)),
)
def grade_many(client: genai.Client, model: str, prompt: str, expected_ids: List[str],
               temperature: float = 1.0, max_output_tokens: int = 8192) -> List[Dict[str, Any]]:
    resp = client.models.generate_content(
        model=model,
        contents=prompt,
        config={
            "temperature": temperature,
            "max_output_tokens": max_output_tokens,
        },
    )
    raw = (resp.text or "").strip()
    if not raw and getattr(resp, "candidates", None):
        parts = []
        for c in resp.candidates:
            if getattr(c, "content", None) and getattr(c.content, "parts", None):
                for p in c.content.parts:
                    if getattr(p, "text", None):
                        parts.append(p.text)
        if parts:
            raw = "\n".join(parts).strip()
    try:
        json_str = extract_json_array(raw)
    except GeminiJSONError:
        debug_dir = Path(__file__).parent / "results"
        debug_dir.mkdir(parents=True, exist_ok=True)
        debug_file = debug_dir / "debug_last_response.txt"
        with debug_file.open("w", encoding="utf-8") as f:
            f.write(f"len(raw)={len(raw)}\n\n--- FULL RAW ---\n{raw}\n")
        print(f"\n[DEBUG] JSON parse failed. Raw ({len(raw)} chars) saved to: {debug_file}")
        print(f"[DEBUG] First 600 chars:\n{(raw[:600] if len(raw) > 600 else raw)}\n")
        raise
    payload = json.loads(json_str)
    graded = validate_list(payload, expected_ids)

    for g in graded:
        g["raw_json"] = json_str
    return graded


# =========================
# Experiment runner
# =========================

def run_grading_one_call(
    client: genai.Client,
    model: str,
    sample_paths: List[Path],
    sample_texts: Dict[str, str],
    base_dir: Path,
    mode: str,   # "fewshot" or "zeroshot"
    order: str,  # "original" or "reversed"
    calib: Optional[Dict[str, str]],
    temperature: float,
) -> pd.DataFrame:
    seq = list(sample_paths)
    if order == "reversed":
        seq = list(reversed(seq))

    # Use path relative to base_dir so same filename in different folders (e.g. Average vs Poor) stay distinct
    ids = [str(p.relative_to(base_dir)) for p in seq]
    essays = [(sid, sample_texts[sid]) for sid in ids]

    if mode == "fewshot":
        prompt = build_fewshot_prompt_multi(
            anchor_exceptional=calib["exceptional"],
            anchor_average=calib["average"],
            anchor_poor=calib["poor"],
            essays=essays,
        )
    else:
        prompt = build_zeroshot_prompt_multi(essays)

    graded = grade_many(
        client=client,
        model=model,
        prompt=prompt,
        expected_ids=ids,
        temperature=temperature,
        max_output_tokens=8192,
    )

    rows = []
    for idx, g in enumerate(graded):
        rows.append({
            "id": g["id"],
            "sample_name": g["id"],
            "position_in_run": idx,
            "mode": mode,
            "order": order,
            "pred_score": g["pred_score"],
            "justification": g["justification"],
            "clamped": g.get("clamped", False),
            "raw_json": g.get("raw_json", ""),
        })
    return pd.DataFrame(rows)


def save_df(df: pd.DataFrame, out_csv: Path, out_jsonl: Path) -> None:
    df.to_csv(out_csv, index=False)
    with out_jsonl.open("w", encoding="utf-8") as f:
        for _, row in df.iterrows():
            f.write(json.dumps(row.to_dict(), ensure_ascii=False) + "\n")


CHECKPOINT_FILENAME = "run_checkpoint.json"


def load_checkpoint(out_dir: Path) -> Optional[Dict[str, Any]]:
    p = out_dir / CHECKPOINT_FILENAME
    if not p.exists():
        return None
    try:
        with p.open("r", encoding="utf-8") as f:
            data = json.load(f)
        if not isinstance(data, dict) or "seed" not in data or "batch_size" not in data or "completed_batches" not in data:
            return None
        return data
    except Exception:
        return None


def save_checkpoint(out_dir: Path, seed: int, batch_size: int, completed_batches: int) -> None:
    with (out_dir / CHECKPOINT_FILENAME).open("w", encoding="utf-8") as f:
        json.dump({"seed": seed, "batch_size": batch_size, "completed_batches": completed_batches}, f, indent=2)


def main():
    load_dotenv()

    api_key = os.getenv("GEMINI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("Missing GEMINI_API_KEY. Put it in your environment or a .env file.")

    model = os.getenv("GEMINI_MODEL", "gemini-3-flash-preview").strip()
    temperature = float(os.getenv("GEMINI_TEMPERATURE", "1.0"))

    # Prefer DATA_BASE_DIR env var; otherwise default to the local final_data_sources folder.
    base_dir = os.getenv("DATA_BASE_DIR", "").strip()
    if not base_dir:
        base_dir = "/Users/arnavpalkhiwala/Desktop/School/UW/STAT 496/STAT496Capstone/final_experiments/final_data_sources"

    print(f"Using model: {model}")
    print(f"Using temperature: {temperature}")
    print(f"Using DATA_BASE_DIR: {base_dir}")

    paths = load_paths(base_dir)
    print(f"Resolved base_dir to: {paths.base_dir}")
    print(f"Output directory will be: {paths.outputs_dir}")

    # Few-shot exemplar filenames (stems)
    fewshot_paths = {
        "exceptional": find_docx_by_stem(paths.exceptional_dir, "exceptional_example"),
        "average":     find_docx_by_stem(paths.average_dir,     "average_example"),
        "poor":        find_docx_by_stem(paths.poor_dir,        "poor_example"),
    }

    calib = {k: docx_to_text(v) for k, v in fewshot_paths.items()}
    print("Loaded calibration exemplars (exceptional, average, poor).")

    # Samples = all .docx in the three folders, excluding the three exemplars
    all_docs = []
    for folder in [paths.exceptional_dir, paths.average_dir, paths.poor_dir]:
        all_docs.extend(sorted(folder.glob("*.docx")))

    exclude = {v.resolve() for v in fewshot_paths.values()}
    sample_paths = [p for p in all_docs if p.resolve() not in exclude]
    print(f"Found {len(sample_paths)} sample essays (excluding exemplars).")

    # Load sample texts keyed by path relative to base_dir (so "Average Essays/essay0.docx" vs "Poor Essays/essay0.docx" are distinct)
    sample_texts = {str(p.relative_to(paths.base_dir)): docx_to_text(p) for p in tqdm(sample_paths, desc="Loading DOCX")}
    print("Finished loading all sample DOCX files into memory.")

    # Randomly shuffle all samples and create non-overlapping batches.
    batch_size = int(os.getenv("BATCH_SIZE", "5"))
    if batch_size <= 0:
        raise RuntimeError("BATCH_SIZE must be a positive integer.")

    out = paths.outputs_dir
    checkpoint = load_checkpoint(out)
    if checkpoint and checkpoint.get("batch_size") == batch_size:
        seed = checkpoint["seed"]
        completed_batches = checkpoint["completed_batches"]
        print(f"Resuming from checkpoint: {completed_batches} batches already done (seed={seed}).")
    else:
        seed = random.randint(0, 2**31 - 1)
        completed_batches = 0
        if checkpoint:
            print("Checkpoint found but batch_size changed; starting fresh.")

    random.seed(seed)
    shuffled_paths = list(sample_paths)
    random.shuffle(shuffled_paths)
    batches = [shuffled_paths[i:i + batch_size] for i in range(0, len(shuffled_paths), batch_size)]
    print(f"Shuffled essays and created {len(batches)} batches (batch size = {batch_size}).")

    client = genai.Client(api_key=api_key)

    for batch_idx in range(completed_batches, len(batches)):
        batch_paths = batches[batch_idx]
        print(f"Processing batch {batch_idx + 1}/{len(batches)} with {len(batch_paths)} essays...")
        df_few_orig = run_grading_one_call(client, model, batch_paths, sample_texts, paths.base_dir, "fewshot", "original", calib, temperature)
        df_few_rev = run_grading_one_call(client, model, batch_paths, sample_texts, paths.base_dir, "fewshot", "reversed", calib, temperature)
        df_zero_orig = run_grading_one_call(client, model, batch_paths, sample_texts, paths.base_dir, "zeroshot", "original", None, temperature)
        df_zero_rev = run_grading_one_call(client, model, batch_paths, sample_texts, paths.base_dir, "zeroshot", "reversed", None, temperature)

        write_header = (batch_idx == 0 and completed_batches == 0)
        csv_mode = "w" if write_header else "a"
        df_few_orig.to_csv(out / "fewshot_original.csv", mode=csv_mode, index=False, header=write_header)
        df_few_rev.to_csv(out / "fewshot_reversed.csv", mode=csv_mode, index=False, header=write_header)
        df_zero_orig.to_csv(out / "zeroshot_original.csv", mode=csv_mode, index=False, header=write_header)
        df_zero_rev.to_csv(out / "zeroshot_reversed.csv", mode=csv_mode, index=False, header=write_header)

        for _df, fname in [
            (df_few_orig, "fewshot_original.jsonl"),
            (df_few_rev, "fewshot_reversed.jsonl"),
            (df_zero_orig, "zeroshot_original.jsonl"),
            (df_zero_rev, "zeroshot_reversed.jsonl"),
        ]:
            with (out / fname).open("a" if not write_header else "w", encoding="utf-8") as f:
                for _, row in _df.iterrows():
                    f.write(json.dumps(row.to_dict(), ensure_ascii=False) + "\n")

        save_checkpoint(out, seed, batch_size, batch_idx + 1)

    print("Completed grading for all batches. Combining results...")

    df_few_orig = pd.read_csv(out / "fewshot_original.csv")
    df_few_rev = pd.read_csv(out / "fewshot_reversed.csv")
    df_zero_orig = pd.read_csv(out / "zeroshot_original.csv")
    df_zero_rev = pd.read_csv(out / "zeroshot_reversed.csv")

    df_all = pd.concat([df_few_orig, df_few_rev, df_zero_orig, df_zero_rev], ignore_index=True)
    save_df(df_all, out/"combined_all_conditions.csv", out/"combined_all_conditions.jsonl")
    print("Saved per-condition and combined results.")

    # scores_by_sample_name.csv (wide format)
    wide = df_all.pivot_table(
        index="sample_name",
        columns=["mode", "order"],
        values="pred_score",
        aggfunc="first"
    ).reset_index()

    # Flatten columns
    wide.columns = ["sample_name"] + [f"{m}_{o}" for (m, o) in wide.columns.tolist()[1:]]
    wide.to_csv(out/"scores_by_sample_name.csv", index=False)
    print("Saved wide-format scores_by_sample_name.csv.")

    print(f"✅ Done. Outputs saved to: {out.resolve()}")
    print("Key files:")
    print(" - combined_all_conditions.csv")
    print(" - scores_by_sample_name.csv")


if __name__ == "__main__":
    main()